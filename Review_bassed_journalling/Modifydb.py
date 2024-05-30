from docx import Document

def remove_line_changes_except_scenario(docx_file):
    doc = Document(docx_file)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    merged_lines = []
    for para in paragraphs:
        if para.startswith("Scenario"):
            merged_lines.append("\n\n" + para)
        else:
            merged_lines.append(para.strip())
    return " ".join(merged_lines)

def replace_multiple_spaces(docx_file):
    doc = Document(docx_file)
    for para in doc.paragraphs:
        text = para.text
        # Split the text by spaces and remove empty strings
        words = [word for word in text.split(' ') if word]
        # Reconstruct the text with a single space between words
        para.text = ' '.join(words)
    doc.save("Scenarios3.docx")

# Example usage:
input_docx_file = "Scenarios.docx"

replace_multiple_spaces(input_docx_file)


